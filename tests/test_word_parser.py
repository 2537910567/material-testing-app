"""Tests for Word (.docx) parser — V4.9"""
import pytest
import sys
import os
import tempfile

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from app.engine.word_parser import extract_word_content


class TestWordParser:
    """Test Word document content extraction."""

    def test_missing_file(self):
        """Non-existent file returns error text."""
        result = extract_word_content("/nonexistent/file.docx")
        assert "text" in result
        assert "ERROR" in result["text"]
        assert result["tables"] == []
        assert result["pages"] == 1

    def test_empty_docx(self):
        """Empty .docx file returns minimal text."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
            doc = Document()
            doc.save(tmp_path)

        try:
            result = extract_word_content(tmp_path)
            assert result["pages"] == 1
            assert isinstance(result["tables"], list)
        finally:
            os.unlink(tmp_path)

    def test_paragraph_extraction(self):
        """Paragraphs are extracted as text."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
            doc = Document()
            doc.add_paragraph("测试段落一：路基工程")
            doc.save(tmp_path)

        try:
            result = extract_word_content(tmp_path)
            assert "路基工程" in result["text"]
        finally:
            os.unlink(tmp_path)

    def test_table_extraction(self):
        """Tables are extracted in Markdown format."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            tmp_path = f.name
            doc = Document()
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "材料"
            table.cell(1, 0).text = "沥青"
            doc.save(tmp_path)

        try:
            result = extract_word_content(tmp_path)
            assert len(result["tables"]) >= 1
        finally:
            os.unlink(tmp_path)
