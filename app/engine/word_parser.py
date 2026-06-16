"""  Word (.docx) 文档解析模块
提取段落文字和表格内容，供 AI 文本分析使用。
V4.9: 新增文件格式支持。
"""

from pathlib import Path
from typing import List, Dict
import re


def extract_word_content(file_path: str) -> dict:
    """
    解析 .docx 文件，提取文字和表格。

    Args:
        file_path: .docx 文件路径

    Returns:
        {"text": str, "tables": List[str], "pages": 1}
        - text: 所有段落文字（保留段落分隔）
        - tables: Markdown 格式的表格文本列表
        - pages: 始终为 1（Word 不分页）
    """
    try:
        from docx import Document
    except ImportError:
        return {"text": "[ERROR: python-docx not installed]", "tables": [], "pages": 1}

    path = Path(file_path)
    if not path.exists():
        return {"text": f"[ERROR: File not found: {file_path}]", "tables": [], "pages": 1}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"text": f"[ERROR: Failed to open {path.name}: {e}]", "tables": [], "pages": 1}

    text_parts = []
    tables_md = []

    # 提取段落
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            if para.style and para.style.name and "heading" in para.style.name.lower():
                text_parts.append(f"\n## {t}")
            else:
                text_parts.append(t)

    # 提取表格
    for ti, table in enumerate(doc.tables):
        rows = []
        max_cols = 0
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            max_cols = max(max_cols, len(cells))
            rows.append(cells)

        if not rows:
            continue

        md_lines = [f"### 表格 {ti + 1}"]
        header = rows[0]
        while len(header) < max_cols:
            header.append("")
        md_lines.append("| " + " | ".join(header) + " |")
        md_lines.append("|" + "|".join(["---"] * max_cols) + "|")

        for row in rows[1:]:
            while len(row) < max_cols:
                row.append("")
            md_lines.append("| " + " | ".join(row) + " |")

        tables_md.append("\n".join(md_lines))
        text_parts.append(f"\n[表格 {ti + 1}]")
        for row in rows:
            text_parts.append(" | ".join(row))

    full_text = "\n".join(text_parts)
    full_text = re.sub(r'\n{3,}', '\n\n', full_text)

    return {
        "text": full_text,
        "tables": tables_md,
        "pages": 1,
    }
